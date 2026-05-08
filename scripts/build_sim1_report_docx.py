#!/usr/bin/env python3

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DOCS_SKILL_DIR = Path(
    "/Users/dyl/.codex/plugins/cache/openai-primary-runtime/documents/26.430.10722/skills/documents/scripts"
)
if str(DOCS_SKILL_DIR) not in sys.path:
    sys.path.insert(0, str(DOCS_SKILL_DIR))

from table_geometry import apply_table_geometry, column_widths_from_weights


DOC_LABEL = "sim1 实验报告"
ACCENT = RGBColor(31, 78, 121)
MUTED = RGBColor(90, 90, 90)
LIGHT_FILL = "EAF2F8"
GRID = "D9E2F3"


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = GRID, size: int = 8) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for side in ("top", "left", "bottom", "right"):
        elem = tc_borders.find(qn(f"w:{side}"))
        if elem is None:
            elem = OxmlElement(f"w:{side}")
            tc_borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), str(size))
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    r = paragraph.add_run()
    r._r.append(begin)
    r._r.append(instr)
    r._r.append(separate)
    r._r.append(text)
    r._r.append(end)
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED

    run = paragraph.add_run(" 页")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, before, after in (
        ("Title", 22, 0, 18),
        ("Heading 1", 16, 12, 6),
        ("Heading 2", 13, 10, 4),
        ("Heading 3", 12, 8, 3),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT if name != "Title" else RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    if "Subtitle" not in doc.styles:
        subtitle = doc.styles.add_style("Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    else:
        subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Arial"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(14)


def setup_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    header_p = section.header.paragraphs[0]
    header_p.text = DOC_LABEL
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(3)
    for run in header_p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED

    p_pr = header_p._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    p_bdr.append(bottom)

    footer_p = section.footer.paragraphs[0]
    add_page_field(footer_p)


def iter_blocks(lines: list[str]) -> Iterable[tuple[str, object]]:
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            yield ("table", table_lines)
            continue

        if re.match(r"^#{1,6}\s+", stripped):
            level = len(stripped) - len(stripped.lstrip("#"))
            yield ("heading", (level, stripped[level:].strip()))
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i].strip()))
                i += 1
            yield ("olist", items)
            continue

        if stripped.startswith("- "):
            items = []
            while i < len(lines) and lines[i].strip().startswith("- "):
                items.append(lines[i].strip()[2:].strip())
                i += 1
            yield ("ulist", items)
            continue

        if stripped.startswith(">"):
            quotes = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quotes.append(lines[i].strip()[1:].strip())
                i += 1
            yield ("quote", " ".join(quotes))
            continue

        paras = [stripped]
        i += 1
        while i < len(lines):
            probe = lines[i].strip()
            if not probe:
                i += 1
                break
            if (
                probe.startswith("|")
                or probe.startswith("- ")
                or probe.startswith(">")
                or re.match(r"^#{1,6}\s+", probe)
                or re.match(r"^\d+\.\s+", probe)
            ):
                break
            paras.append(probe)
            i += 1
        yield ("para", " ".join(paras))


def clean_text(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.replace("`", "")


def add_md_table(doc: Document, table_lines: list[str]) -> None:
    rows = []
    for idx, line in enumerate(table_lines):
        parts = [clean_text(p.strip()) for p in line.strip("|").split("|")]
        if idx == 1 and all(re.fullmatch(r":?-{3,}:?", p) for p in parts):
            continue
        rows.append(parts)
    if not rows:
        return

    cols = max(len(r) for r in rows)
    word_table = doc.add_table(rows=len(rows), cols=cols, style="Table Grid")
    word_table.autofit = False

    raw_widths = [max(len(row[c]) if c < len(row) else 0 for row in rows) for c in range(cols)]
    weights = [max(w, 8) for w in raw_widths]
    widths = column_widths_from_weights(weights, total_width_dxa=9360)

    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = word_table.cell(r_idx, c_idx)
            cell.text = row[c_idx] if c_idx < len(row) else ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(3)
                if r_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif c_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9.5)
                    if r_idx == 0:
                        run.font.bold = True
            if r_idx == 0:
                set_cell_fill(cell, LIGHT_FILL)
            set_cell_border(cell)

    apply_table_geometry(word_table, widths, table_width_dxa=sum(widths))
    doc.add_paragraph("")


def build_doc(input_path: Path, output_path: Path) -> None:
    doc = Document()
    configure_styles(doc)
    setup_page(doc)

    lines = input_path.read_text(encoding="utf-8").splitlines()
    title = "sim1 实验报告"
    subtitle = "由 Markdown 自动整理为 Word 文档"

    # Consume the first title line, if present.
    if lines and lines[0].startswith("# "):
        title = lines[0][2:].strip()
        lines = lines[1:]

    p = doc.add_paragraph(title, style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle_p = doc.add_paragraph(subtitle, style="Subtitle")
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_table(rows=2, cols=2, style="Table Grid")
    meta.cell(0, 0).text = "来源"
    meta.cell(0, 1).text = str(input_path)
    meta.cell(1, 0).text = "说明"
    meta.cell(1, 1).text = "本版用于论文写作与汇报，可直接在 Word 中继续修改。"
    meta_widths = [1400, 7960]
    for r_idx, row in enumerate(meta.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if c_idx == 0:
                set_cell_fill(cell, LIGHT_FILL)
            set_cell_border(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
                    if c_idx == 0:
                        run.font.bold = True
    apply_table_geometry(meta, meta_widths, table_width_dxa=sum(meta_widths))
    doc.add_paragraph("")

    ordered_counter = 1
    for kind, payload in iter_blocks(lines):
        if kind == "heading":
            level, text = payload
            style = {2: "Heading 1", 3: "Heading 2", 4: "Heading 3"}.get(level, "Heading 1")
            doc.add_paragraph(clean_text(text), style=style)
            ordered_counter = 1
        elif kind == "para":
            p = doc.add_paragraph(style="Normal")
            p.add_run(clean_text(payload))
        elif kind == "ulist":
            for item in payload:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(clean_text(item))
        elif kind == "olist":
            for item in payload:
                p = doc.add_paragraph(style="Normal")
                p.paragraph_format.left_indent = Inches(0.15)
                p.paragraph_format.first_line_indent = Inches(-0.15)
                p.add_run(f"{ordered_counter}. {clean_text(item)}")
                ordered_counter += 1
        elif kind == "quote":
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(clean_text(payload))
            run.italic = True
            run.font.color.rgb = RGBColor(70, 70, 70)
        elif kind == "table":
            add_md_table(doc, payload)
            ordered_counter = 1

    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input_md", type=Path)
    parser.add_argument("output_docx", type=Path)
    args = parser.parse_args()
    build_doc(args.input_md, args.output_docx)


if __name__ == "__main__":
    main()
